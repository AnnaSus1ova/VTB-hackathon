# This Python 3 environment comes with many helpful analytics libraries installed
# It is defined by the kaggle/python Docker image: https://github.com/kaggle/docker-python
# For example, here's several helpful packages to load

import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
import os, io, re
import torch
import pandas as pd
import torch.nn.functional as F
from transformers import pipeline, AutoTokenizer, AutoModel, AutoModelForSequenceClassification
from docx import Document
from striprtf.striprtf import rtf_to_text
import mammoth

# Input data files are available in the read-only "../input/" directory
# For example, running this (by clicking run or pressing Shift+Enter) will list all files under the input directory

import os
for dirname, _, filenames in os.walk('/kaggle/input'):
    for filename in filenames:
        print(os.path.join(dirname, filename))

# You can write up to 20GB to the current directory (/kaggle/working/) that gets preserved as output when you create a version using "Save & Run All" 
# You can also write temporary files to /kaggle/temp/, but they won't be saved outside of the current session
device = "cuda" if torch.cuda.is_available() else "cpu"

DATASET_PATH = "./VTB hackaton/kaggle/input/for-llama"
OUT_CSV = "./VTB hackaton/kaggle/working/resume_vacancy_match_improved_fixed.csv"
TARGET_DIM = 512 

VEC_DIR = "./candidate_vectors"  # Папка для векторов кандидатов
SUMMARY_DIR = "./candidates_summary"  # Папка для summary

# Создайте папки
os.makedirs(VEC_DIR, exist_ok=True)
os.makedirs(SUMMARY_DIR, exist_ok=True)

def read_docx_paragraphs(path):
    try:
        doc = Document(path)
        texts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                texts.append(p.text.strip())
        # также соберём текст из таблиц, если есть
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        texts.append(cell.text.strip())
        return "\n".join(texts).strip()
    except Exception as e:
        return ""

def read_docx_mammoth(path):
    try:
        with open(path, "rb") as f:
            res = mammoth.extract_raw_text(f)
            txt = res.value or ""
            return txt.strip()
    except Exception:
        return ""

def read_rtf_striprtf(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        txt = rtf_to_text(raw)
        return txt.strip()
    except Exception:
        return ""

def read_rtf_fallback(path):
    # простая очистка: удалить управляющие последовательности { ... } и \xxx
    try:
        with open(path, "rb") as f:
            rawb = f.read()
        # попробуем декодить cp1251 (русский), затем utf-8
        for enc in ("utf-8","cp1251","latin1"):
            try:
                raw = rawb.decode(enc)
                break
            except Exception:
                raw = None
        if raw is None:
            raw = rawb.decode("utf-8", errors="ignore")
        # удалить rtf-метки
        cleaned = re.sub(r"[\\][a-zA-Z]+\d*", " ", raw)   # \par \b0 etc
        cleaned = re.sub(r"\{|\}", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip()
    except Exception:
        return ""

def read_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        # 1) try paragraphs+tables (python-docx)
        t = read_docx_paragraphs(path)
        if t:
            return t
        # 2) try mammoth (fallback)
        t = read_docx_mammoth(path)
        if t:
            return t
        # 3) last resort: raw binary read + try cp1251/utf8
        try:
            with open(path, "rb") as f:
                rawb = f.read()
            for enc in ("utf-8","cp1251","latin1"):
                try:
                    txt = rawb.decode(enc)
                    if txt and any(c.isalpha() for c in txt[:200]):
                        return txt.strip()
                except Exception:
                    pass
        except Exception:
            pass
        return ""
    elif ext == ".rtf":
        t = read_rtf_striprtf(path)
        if t:
            return t
        t = read_rtf_fallback(path)
        return t or ""
    else:
        # try plain text decode
        try:
            with open(path, "rb") as f:
                rawb = f.read()
            for enc in ("utf-8","cp1251","latin1"):
                try:
                    txt = rawb.decode(enc)
                    if txt and any(c.isalpha() for c in txt[:200]):
                        return txt.strip()
                except Exception:
                    pass
        except Exception:
            pass
        return ""

# -------------------------
# Собираем descriptions и resumes, логируем размеры
# -------------------------
descriptions = {}
resumes = {}
stats = []

for fn in sorted(os.listdir(DATASET_PATH)):
    if not fn.lower().startswith(("description","resume")):
        continue
    path = os.path.join(DATASET_PATH, fn)
    if not os.path.isfile(path):
        continue
    text = read_file(path)
    stats.append((fn, len(text)))
    if fn.lower().startswith("description"):
        descriptions[fn] = text
    else:
        resumes[fn] = text

print("Найдено описаний вакансий:", list(descriptions.keys()))
print("Найдено резюме:", list(resumes.keys()))
print("Длины (символов):")
for fn, ln in stats:
    print(f"  {fn}: {ln}")


for fn, txt in list(descriptions.items())[:2]:
    print("=== DESC SAMPLE", fn, "===")
    print(txt[:400])
for fn, txt in list(resumes.items())[:4]:
    print("=== RESUME SAMPLE", fn, "===")
    print(txt[:400])

from transformers import pipeline

device = 0 if torch.cuda.is_available() else -1
clf = pipeline("zero-shot-classification",
               model="joeddav/xlm-roberta-large-xnli",
               device=device)

print(clf("Я люблю машинное обучение", candidate_labels=["спорт", "наука", "музыка"]))

# Подготовка NLI и эмбеддингов (русский/мультиязычный)
# -------------------------

EMB_MODEL_NAME = "intfloat/multilingual-e5-small"
emb_tokenizer = AutoTokenizer.from_pretrained(EMB_MODEL_NAME)
emb_model = AutoModel.from_pretrained(EMB_MODEL_NAME).to("cuda" if torch.cuda.is_available() else "cpu")
emb_model.eval()

def get_emb(text, max_length=512):
    text = text.strip() if text else ""
    if not text:
        return torch.zeros(emb_model.config.hidden_size, device=emb_model.device)
    inputs = emb_tokenizer("passage: " + text, return_tensors="pt", truncation=True, padding="max_length", max_length=max_length).to(emb_model.device)
    with torch.no_grad():
        out = emb_model(**inputs)
        last = out.last_hidden_state
        mask = inputs["attention_mask"].unsqueeze(-1)
        summed = (last * mask).sum(1)
        denom = mask.sum(1).clamp(min=1e-9)
        emb = summed / denom
        emb = F.normalize(emb, p=2, dim=1)
    return emb.squeeze(0)


def normalize_meta_val(x, low=0.0, high=1.0):
    x = float(x)
    return max(0.0, min(1.0, (x - low) / (high - low + 1e-12)))
    

def build_candidate_vector(resume_text, desc_text=None, summary_text=None,
                           label_confidence=0.0, similarity=0.0, combined=0.0,
                           emb_model=get_emb, target_dim=TARGET_DIM):
    
    # Определяем устройство
    current_device = "cuda" if torch.cuda.is_available() else "cpu"
    
    # Получаем эмбеддинги и явно указываем устройство
    r_emb = emb_model(resume_text).to(current_device)
    d_emb = emb_model(desc_text).to(current_device) if desc_text else torch.zeros_like(r_emb).to(current_device)
    s_emb = emb_model(summary_text).to(current_device) if summary_text else torch.zeros_like(r_emb).to(current_device)

    # Веса для комбинирования
    w_r, w_d, w_s = 0.6, 0.2, 0.2
    combined_emb = w_r * r_emb + w_d * d_emb + w_s * s_emb

    # Мета-данные
    meta = torch.tensor([
        normalize_meta_val(label_confidence, 0.0, 1.0),
        normalize_meta_val(similarity, 0.0, 1.0),
        normalize_meta_val(combined, 0.0, 1.0)
    ], device=current_device, dtype=combined_emb.dtype)

    # Повторяем мета-данные для совместимости размеров
    meta_repeated = meta.repeat(combined_emb.shape[0] // len(meta) + 1)[:combined_emb.shape[0]]
    full = torch.cat([combined_emb, meta_repeated], dim=0)

    # Проекция в target_dim
    full = full.unsqueeze(0)
    N = full.shape[1]
    
    if N > target_dim:
        try:
            M = full - full.mean(dim=1, keepdim=True)
            U, S, Vt = torch.linalg.svd(M, full_matrices=False)
            proj = (M @ Vt.T[:, :target_dim]).squeeze(0)
        except Exception:
            proj = M.squeeze(0)[:target_dim]
    else:
        proj = torch.zeros(target_dim, device=current_device, dtype=full.dtype)
        proj[:N] = full.squeeze(0)

    # Нормализация
    proj = proj / (proj.norm(p=2) + 1e-9)

    return proj.cpu().numpy().astype(np.float32)

# -------------------------
# Основная логика: classify -> similarity
# -------------------------

def label_from_desc(fn, text):
    if text:
        for l in text.splitlines():
            l = l.strip()
            # Игнорируем короткие, шаблонные или служебные строки
            if len(l) > 3 and not any(bad in l.lower() for bad in ["наименование", "значение", "field", "column"]):
                return l.split("|")[0].split("-")[0][:60].strip()
    # fallback по имени файла
    low = fn.lower()
    if "it" in low or "специал" in low:
        return "Специалист IT"
    if "analit" in low or "аналит" in low:
        return "Бизнес-аналитик"
    return os.path.splitext(fn)[1]

desc_labels = {fn: label_from_desc(fn, txt) for fn, txt in descriptions.items()}
labels = sorted(set(desc_labels.values()))
print("Derived labels:", labels)

THRESH_CLASS = 0.45
results = []
for r_fn, r_text in resumes.items():
    if not r_text or not r_text.strip():
        print("Empty resume text (after parsing):", r_fn)
        results.append({
            "Predicted_Label": None,
            "Label_Confidence": 0.0,
            "Vacancy_File": None,
            "Resume_File": r_fn,
            "Similarity": 0.0,
            "Combined_Score": 0.0
        })
        continue

    try:
        pred = clf(r_text, candidate_labels=labels, hypothesis_template="Это резюме для роли: {}.")
    except Exception as e:
        print("Zero-shot error for", r_fn, ":", e)
        pred = {"labels": [], "scores": []}

    if not pred["labels"]:
        best_label = None
        best_score = 0.0
    else:
        best_label = pred["labels"][0]
        best_score = float(pred["scores"][0])

    if best_label is None or best_score < THRESH_CLASS:
        # не отнесено ни к одной роли уверенно
        results.append({
            "Predicted_Label": None,
            "Label_Confidence": round(best_score,4),
            "Vacancy_File": None,
            "Resume_File": r_fn,
            "Similarity": 0.0,
            "Combined_Score": 0.0
        })
        continue

    # candidate description files
    candidate_desc_files = [fn for fn, lbl in desc_labels.items() if lbl == best_label]
    if not candidate_desc_files:
        results.append({
            "Predicted_Label": best_label,
            "Label_Confidence": round(best_score,4),
            "Vacancy_File": None,
            "Resume_File": r_fn,
            "Similarity": 0.0,
            "Combined_Score": 0.0
        })
        continue

    r_emb = get_emb(r_text)
    best_sim = -1.0
    best_desc = None
    for d_fn in candidate_desc_files:
        d_text = descriptions.get(d_fn, "")
        d_emb = get_emb(d_text)
        sim = float(torch.dot(r_emb, d_emb).cpu().item())
        if sim > best_sim:
            best_sim = sim
            best_desc = d_fn

    combined = best_sim * best_score
    results.append({
        "Predicted_Label": best_label,
        "Label_Confidence": round(best_score,4),
        "Vacancy_File": best_desc,
        "Resume_File": r_fn,
        "Similarity": round(best_sim,4),
        "Combined_Score": round(combined,4)
    })
# -------------------------
# Сохраняем и печатаем
# -------------------------
df = pd.DataFrame(results)
df = df.sort_values(by="Combined_Score", ascending=False)
df.to_csv(OUT_CSV, index=False)
print("Saved CSV:", OUT_CSV)
print(df)

# --- Summarizer ---
SUM_MODEL = "cointegrated/rut5-base-multitask"
summarizer = pipeline(
    "summarization",
    model=SUM_MODEL,
    tokenizer=SUM_MODEL,
    device=device
)

import re
from transformers import pipeline
THRESH_PASS = 0.6
SUMMARY_DIR = "candidates_summary"
os.makedirs(SUMMARY_DIR, exist_ok=True)

# --- Очистка текста ---
def clean_text(text: str) -> str:
    # убираем все странные символы и мусор от docx/rtf
    text = text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
    text = re.sub(r"[^а-яА-Яa-zA-Z0-9\s.,;:!?()-]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text



# --- Добавляем колонку для итогового summary ---
df["Candidate_Summary"] = ""

# --- Цикл по кандидатам ---
for idx, row in df.iterrows():
    if row["Combined_Score"] > THRESH_PASS:
        print(f"\n✅ {row['Resume_File']} проходит на собеседование (score={row['Combined_Score']:.4f})")

        resume_text = resumes.get(row["Resume_File"], "")
        resume_text = clean_text(resume_text)

        summary = ""
        if resume_text:
            try:
                # Добавляем инструкцию для T5
                summary = summarizer(
                    "summarize: " + resume_text[:2500],  # меньше текста на вход
                    max_length=200,
                    min_length=50,
                    repetition_penalty=2.5,   # штраф за повторы
                    num_beams=5,              # beam search
                    no_repeat_ngram_size=3,   # запрет на повтор триграмм
                    do_sample=False
                )[0]["summary_text"]
            except Exception as e:
                print("Ошибка суммаризации:", e)
                summary = "\n".join(resume_text.splitlines()[:5])

        # Сохраняем и печатаем
        # Сохраняем в DataFrame
        df.at[idx, "Candidate_Summary"] = summary

        # Сохраняем в отдельный файл
        safe_name = os.path.splitext(os.path.basename(row["Resume_File"]))[0]
        out_path = os.path.join(SUMMARY_DIR, f"{safe_name}_summary.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(f"Файл резюме: {row['Resume_File']}\n")
            f.write(f"Score: {row['Combined_Score']:.4f}\n\n")
            f.write("Представление о кандидате:\n")
            f.write(summary.strip() + "\n")

        print("Представление о кандидате:")
        print(summary)
        
# --- Строим и сохраняем векторы ---
df["Vector_Path"] = ""

for idx, row in df.iterrows():
    resume_text = resumes.get(row["Resume_File"], "")
    desc_text = descriptions.get(row["Vacancy_File"], "")
    summary_text = row.get("Candidate_Summary", "")

    vec = build_candidate_vector(
        resume_text=resume_text,
        desc_text=desc_text,
        summary_text=summary_text,
        label_confidence=row["Label_Confidence"],
        similarity=row["Similarity"],
        combined=row["Combined_Score"]
    )

    vec_path = os.path.join(VEC_DIR, row["Resume_File"] + ".npy")
    np.save(vec_path, vec)
    df.at[idx, "Vector_Path"] = vec_path
# --- Сохраняем в CSV ---
df.to_csv("results_with_summary.csv", index=False, encoding="utf-8-sig")
print("\n📄 Итог сохранён в results_with_summary.csv")


